import os
import argparse
from docx import Document

# Function to replace text in a docx file


def replace_text_in_docx(doc_path, search_text, replace_text):
    try:
        doc = Document(doc_path)

        # Loop through paragraphs in the document
        for para in doc.paragraphs:
            if search_text in para.text:
                para.text = para.text.replace(search_text, replace_text)

        # Loop through tables in the document, if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_text in cell.text:
                        cell.text = cell.text.replace(
                            search_text, replace_text)

        # Save the modified document
        doc.save(doc_path)

    except Exception as e:
        # Log the error and file path to error.txt
        with open("error.txt", "a") as error_log:
            error_log.write(f"Failed to process {doc_path}: {e}\n")


# Main function to handle command-line arguments and process files


def main():
    # Set up command-line argument parsing
    parser = argparse.ArgumentParser(
        description="Replace text in .docx files recursively."
    )

    # Add arguments
    parser.add_argument(
        "folder_path", type=str, help="Path to the folder containing .docx files"
    )
    parser.add_argument("word_to_replace", type=str,
                        help="The word to be replaced")
    parser.add_argument("replace_by", type=str,
                        help="The word to replace with")

    # Parse the arguments
    args = parser.parse_args()

    # Recursively walk through the folder and all subfolders
    for root, dirs, files in os.walk(args.folder_path):
        for filename in files:
            if filename.endswith(".docx"):
                file_path = os.path.join(root, filename)
                replace_text_in_docx(
                    file_path, args.word_to_replace, args.replace_by)
                print(f"Processed: {file_path}")


if __name__ == "__main__":
    main()
